import os
import fitz  # PyMuPDF
import docx
import requests
import pandas as pd
from config import ANTHROPIC_API_KEY
import re

def read_pdf(filepath):
    """Extract text from PDF with improved formatting and structure"""
    text = ""
    try:
        with fitz.open(filepath) as doc:
            # First pass: Get detailed structure
            for page_num, page in enumerate(doc):
                print(f"\nProcessing page {page_num + 1}")
                
                # Try multiple extraction methods
                text_dict = page.get_text("dict", sort=True)
                page_text = []
                
                # Extract text with hierarchy
                for block in text_dict["blocks"]:
                    if "lines" in block:
                        block_text = []
                        for line in block["lines"]:
                            # Preserve formatting and structure
                            spans = []
                            for span in line["spans"]:
                                span_text = span["text"].strip()
                                if span_text:
                                    # Check for potential section headers or important data
                                    if span.get("size", 0) > 10:  # Larger text might be headers
                                        span_text = f"\n{span_text}\n"
                                    spans.append(span_text)
                            
                            line_text = " ".join(spans)
                            if line_text.strip():
                                block_text.append(line_text)
                        
                        if block_text:
                            block_content = "\n".join(block_text)
                            # Preserve paragraph breaks and structure
                            if block.get("type", 0) == 0:  # Text block
                                page_text.append(block_content)
                            else:  # Other block types (images, tables, etc.)
                                page_text.append(f"\n{block_content}\n")
                
                text += "\n".join(page_text) + "\n\n"
            
            # If structured extraction failed, try raw text
            if not text.strip():
                print("Structured extraction failed, trying raw text...")
                text = ""
                for page in doc:
                    text += page.get_text("text") + "\n\n"
            
            # Clean up the text
            text = text.replace("\n\n\n", "\n\n")
            text = text.strip()
            
            # Extract key information for validation
            # Try multiple patterns for tenant name
            tenant_patterns = [
                r'(?:called\s*["\']Tenant["\']\s*,?\s*)(.*?)(?:\s*[,\n])',
                r'(?:TENANT:\s*)(.*?)(?:\s*[,\n])',
                r'(?:Tenant[:\s]+)(.*?)(?:\s*[,\n])',
                r'(?:between.*?and\s+)(.*?)(?:\s*(?:,|\(|as))',
            ]
            
            tenant_name = None
            for pattern in tenant_patterns:
                tenant_match = re.search(pattern, text, re.IGNORECASE)
                if tenant_match:
                    tenant_name = tenant_match.group(1).strip()
                    break
            
            # Find square footage with improved pattern
            footage_pattern = r'(\d{1,6}(?:,\d{3})*)\s*(?:square\s*feet|sq\.?\s*ft\.?|SF)'
            footage_match = re.search(footage_pattern, text, re.IGNORECASE)
            square_footage = footage_match.group(1) if footage_match else None
            
            print("\nExtracted key information:")
            if tenant_name:
                print(f"- Tenant: {tenant_name}")
            if square_footage:
                print(f"- Square Footage: {square_footage}")
            
            print(f"\nExtracted text statistics:")
            print(f"- Total length: {len(text)} characters")
            print(f"- Number of paragraphs: {text.count('\n\n') + 1}")
            print(f"- First 500 chars: {text[:500]}")
            print(f"- Last 500 chars: {text[-500:]}")
            
            return text, tenant_name, square_footage
            
    except Exception as e:
        print(f"Error reading PDF: {e}")
        raise

def read_docx(filepath):
    """Extract text from DOCX with improved formatting"""
    try:
        doc = docx.Document(filepath)
        text = []
        
        # Process paragraphs with better formatting
        for para in doc.paragraphs:
            if para.text.strip():
                # Preserve paragraph formatting
                text.append(para.text.strip())
        
        # Process tables if present
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text.append(" | ".join(row_text))
        
        return "\n\n".join(text)
    except Exception as e:
        print(f"Error reading DOCX: {e}")
        raise

def generate_prompt(text, tenant_name=None, square_footage=None):
    """Generate prompt with document-specific validation"""
    # Build list of verified information
    verified_info = []
    if tenant_name:
        verified_info.append(f"- Tenant Name: {tenant_name}")
    if square_footage:
        verified_info.append(f"- Square Footage: {square_footage} square feet")
    
    prompt = f"""
You are an expert lease analyst. Analyze the following commercial lease document and extract ONLY THE ACTUAL INFORMATION FOUND IN THE DOCUMENT. Do not use placeholders or generic data.

{f'''IMPORTANT: The following information has been verified to exist in the document:
{chr(10).join(verified_info)}

YOU MUST INCLUDE THIS INFORMATION IN YOUR ANALYSIS. Do not mark these items as N/A as they are confirmed to exist in the text.
''' if verified_info else ''}

Format the output EXACTLY as shown below, maintaining the exact structure and headers:

Section 1: Key Lease Information
Create a table with these EXACT columns: Item | Paragraph Reference | Information
Extract ONLY information that exists in the document for these fields:
- Property Name (exactly as written in document)
- Property Address (full address as shown)
- Property ID/Number (if specified)
- Tenant Name (exact legal name as written)
- Square Footage (exact number from document)
- Amendment Type (exact type mentioned)
- Lease Type (as specified in document)
- Late Fee Details (exact amounts and terms)
- Security Deposit Amount (exact amount)
- Lease Holdover Amount/Rate (exact terms)
- Tenant Contact Information (exact details provided)
- All Dates (exactly as specified in document)
- Suite Details (exact numbers and details provided)

Section 2: Lease Charges
Create a table with these EXACT columns: Charge Type | Frequency | Start Date | End Date | Amount
Include ONLY charges explicitly mentioned in the document:
- Base Rent amounts (exact figures)
- Additional charges (exactly as specified)
- Use exact dates and amounts from the document
- Format dates as MM/DD/YYYY
- Include dollar signs and commas in amounts

Section 3: Lease Options
Create a table with these EXACT columns: Option Type | Expiration Date | Latest Notice | Earliest Notice | Notice to Tenant | Reference
Include ONLY options explicitly stated in the document:
- List each option with its exact terms
- Use actual section references
- Include all notice periods as specified
- Use N/A only when a field is not specified in the document

Section 4: Lease Clauses
Create a table with these EXACT columns: Clause Name | Section/Paragraph | Summary of Clause
Include ONLY clauses found in the document:
- List significant clauses with their exact section numbers
- Provide concise but accurate summaries of each clause
- Include key terms and conditions as written
- Use actual section numbers from the document

IMPORTANT RULES:
1. Extract ONLY information that actually exists in the document
2. Do not generate, assume, or use placeholder data
3. Use 'N/A' only when information is truly not present
4. Maintain exact numbers, dates, and terms as written
5. Include section/paragraph references where available
6. Keep summaries accurate and based on actual text

Here is the document to analyze:

{text}

VERIFICATION REQUIRED: Your response MUST include any verified information listed above. These items are confirmed to exist in the document and should not be marked as N/A."""
    return prompt

def send_to_api(prompt):
    url = "https://api.anthropic.com/v1/messages"
    headers = {
        "x-api-key": ANTHROPIC_API_KEY,
        "anthropic-version": "2023-06-01",
        "content-type": "application/json"
    }
    
    data = {
        "model": "claude-3-opus-20240229",
        "messages": [
            {
                "role": "system",
                "content": prompt
            }
        ],
        "max_tokens": 4000,
        "temperature": 0.1,
        "system": """You are an expert lease document analyzer. Your task is to extract ONLY factual information that exists in the document. 
Never generate, assume, or use placeholder data. If information is not found, mark it as 'N/A'.
IMPORTANT: When specific information is provided as "verified to exist", you MUST include this information in your analysis.
Double check your response before sending to ensure all verified information is included."""
    }
    
    try:
        response = requests.post(url, json=data, headers=headers)
        response.raise_for_status()
        
        # Validate response contains verified information
        response_json = response.json()
        response_text = response_json.get("content", [{}])[0].get("text", "")
        
        if "Wink Davis Equipment Company" not in response_text:
            print("Warning: API response missing verified tenant name")
        if "12,085" not in response_text:
            print("Warning: API response missing verified square footage")
        
        return response_json
    except requests.exceptions.HTTPError as http_err:
        print(f"HTTP error occurred: {http_err}")
        if response.status_code == 404:
            print("Check API endpoint and version")
        elif response.status_code == 401:
            print("Check API key")
        raise
    except Exception as err:
        print(f"Other error occurred: {err}")
        raise

def parse_api_response(response_text):
    """Parse the API response into structured data with improved accuracy"""
    try:
        print("\nStarting to parse API response...")
        print("Full response length:", len(response_text))
        print("Response preview:", response_text[:500])
        
        sections = response_text.split('\n\n')
        tables = {
            'key_info': [],
            'charges': [],
            'options': [],
            'clauses': []
        }
        
        current_section = None
        headers_seen = False
        
        for section in sections:
            section = section.strip()
            if not section:
                continue
            
            # Identify sections
            if 'Section 1:' in section:
                current_section = 'key_info'
                headers_seen = False
            elif 'Section 2:' in section:
                current_section = 'charges'
                headers_seen = False
            elif 'Section 3:' in section:
                current_section = 'options'
                headers_seen = False
            elif 'Section 4:' in section:
                current_section = 'clauses'
                headers_seen = False
            
            if current_section and '|' in section:
                rows = [row.strip() for row in section.split('\n') if row.strip() and '|' in row]
                
                for row in rows:
                    if not headers_seen:
                        headers_seen = True
                        continue
                    
                    if '---' in row:
                        continue
                        
                    cells = [cell.strip() for cell in row.split('|') if cell.strip()]
                    if cells:
                        # Only skip if ALL cells are N/A or empty
                        if not all(cell in ['N/A', ''] for cell in cells):
                            tables[current_section].append(cells)
                            print(f"Added to {current_section}: {cells}")
        
        return tables
    except Exception as e:
        print(f"Error parsing API response: {e}")
        print(f"Response text: {response_text[:500]}")
        return {
            'key_info': [],
            'charges': [],
            'options': [],
            'clauses': []
        }

def convert_to_csv(tables):
    """Convert parsed tables to CSV format"""
    csv_data = []
    
    # Section 1: Key Lease Information
    if tables['key_info']:
        csv_data.append(['Section 1: Key Lease Information'])
        csv_data.append(['Item', 'Paragraph Reference', 'Information'])
        csv_data.extend(tables['key_info'])
        csv_data.append([])  # Empty row for separation
    
    # Section 2: Lease Charges
    if tables['charges']:
        csv_data.append(['Section 2: Lease Charges'])
        csv_data.append(['Charge Type', 'Frequency', 'Start Date', 'End Date', 'Amount'])
        csv_data.extend(tables['charges'])
        csv_data.append([])
    
    # Section 3: Lease Options
    if tables['options']:
        csv_data.append(['Section 3: Lease Options'])
        csv_data.append(['Option Type', 'Expiration Date', 'Latest Notice', 'Earliest Notice', 'Notice to Tenant', 'Reference'])
        csv_data.extend(tables['options'])
    
    # Section 4: Lease Clauses
    if tables['clauses']:
        csv_data.append(['Section 4: Lease Clauses'])
        csv_data.append(['Clause Name', 'Section/Paragraph', 'Summary of Clause'])
        csv_data.extend(tables['clauses'])
    
    # Convert to CSV
    if not csv_data:
        return "No data available"
    
    df = pd.DataFrame(csv_data)
    return df.to_csv(index=False, header=False)

def process_document(filepath):
    ext = filepath.rsplit(".", 1)[1].lower()
    try:
        if ext == "pdf":
            text, tenant_name, square_footage = read_pdf(filepath)
        elif ext == "docx":
            text = read_docx(filepath)
            tenant_name = None
            square_footage = None
        else:
            raise ValueError("Unsupported file type")
        
        print("\n=== EXTRACTED TEXT ===")
        print(text[:1000])
        print("=== END EXTRACTED TEXT ===\n")
        
        prompt = generate_prompt(text, tenant_name, square_footage)
        api_response = send_to_api(prompt)
        
        print("\n=== API RESPONSE ===")
        extracted_data = api_response.get("content")[0].get("text")
        print(extracted_data[:1000])
        print("=== END API RESPONSE ===\n")
        
        # Validate API response includes verified information
        if tenant_name and tenant_name not in extracted_data:
            print(f"Warning: API response missing verified tenant name: {tenant_name}")
        if square_footage and square_footage not in extracted_data:
            print(f"Warning: API response missing verified square footage: {square_footage}")
        
        parsed_data = parse_api_response(extracted_data)
        
        print("\n=== PARSED DATA ===")
        for section, data in parsed_data.items():
            print(f"{section}: {len(data)} entries")
            if data:
                print(f"First entry: {data[0]}")
        print("=== END PARSED DATA ===\n")
        
        csv_data = convert_to_csv(parsed_data)
        return extracted_data, csv_data
    except Exception as e:
        print(f"Error in process_document: {e}")
        raise